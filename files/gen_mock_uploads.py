"""
saasmetrics.ai  |  Generate mock upload files for demo
Run: python gen_mock_uploads.py
Output: mock_uploads/ folder with 3 files ready to drag-and-drop during demo
"""

import json
from pathlib import Path
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, black, white
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent / "mock_uploads"
OUT.mkdir(exist_ok=True)


# ════════════════════════════════════════════════════════════════
# MOCK FILE 1: Competitor Win/Loss Analysis (Excel)
# "I have this from our sales team — let me upload it"
# Will be queried against BQ pipeline data
# ════════════════════════════════════════════════════════════════
def gen_excel():
    wb = openpyxl.Workbook()

    DARK   = "1F3864"
    MID    = "2E5FA3"
    LIGHT  = "D6E4F0"
    GREEN  = "D9EAD3"
    RED    = "FCE5CD"
    GOLD   = "FFF2CC"

    def hdr(ws, col, row, text):
        c = ws.cell(row=row, column=col, value=text)
        c.font = Font(bold=True, color="FFFFFF", size=10)
        c.fill = PatternFill("solid", fgColor=MID)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.border = Border(
            bottom=Side(style="thin", color="FFFFFF"),
            right=Side(style="thin", color="444444"),
        )
        return c

    def cell(ws, col, row, value, bg=None, bold=False, fmt=None):
        c = ws.cell(row=row, column=col, value=value)
        if bg:
            c.fill = PatternFill("solid", fgColor=bg)
        c.font = Font(bold=bold, size=10)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        if fmt:
            c.number_format = fmt
        return c

    # ── Sheet 1: Win/Loss Summary ─────────────────────────────────────────
    ws1 = wb.active
    ws1.title = "Win_Loss_Q4"
    ws1.freeze_panes = "A3"

    ws1.merge_cells("A1:J1")
    title = ws1["A1"]
    title.value = "Competitor Win / Loss Analysis  —  Q4 FY2024"
    title.font = Font(bold=True, color="FFFFFF", size=13)
    title.fill = PatternFill("solid", fgColor=DARK)
    title.alignment = Alignment(horizontal="center", vertical="center")
    ws1.row_dimensions[1].height = 30

    headers = ["Deal", "Account", "ACV ($)", "Competitor", "Outcome", "Primary Reason",
               "Decision Maker", "Sales Cycle (days)", "Rep", "Notes"]
    for i, h in enumerate(headers, 1):
        hdr(ws1, i, 2, h)
    ws1.row_dimensions[2].height = 22

    rows = [
        ("D-201","Fortress Bank NA",       520000,"CrowdStrike","WON",  "Feature parity + 15% price advantage","CISO",        87, "James Okoye",   "RIIP promotion applied"),
        ("D-202","EuroCredit AG",           410000,"SentinelOne","PENDING","Eval in progress — EMEA budget freeze","CTO",        62, "James Okoye",   "Decision deferred Q1 FY25"),
        ("D-203","DataVault Reinsurance",   290000,"Darktrace",  "WON",  "Superior SIEM integration story",     "VP Security", 44, "Priya Nair",    "Proof of concept won it"),
        ("D-204","NexGen Insurance",         96000,"CrowdStrike","LOST", "Price — CrowdStrike 22% cheaper",     "CFO",         31, "Wei Zhang",     "Need to revisit pricing for MM Ins."),
        ("D-205","SwiftTrade LLC",           84000,"Darktrace",  "WON",  "APAC support coverage advantage",     "IT Director", 55, "Wei Zhang",     "24/7 APAC support closed it"),
        ("D-206","Brightpath Credit",        72000,"SentinelOne","LOST", "Incumbent — renewal not displaced",   "CTO",         28, "Priya Nair",    "Need stronger displacement playbook"),
        ("D-207","Pacific Rim Capital",     144000,"Palo Alto",  "WON",  "TI Add-on differentiated the pitch", "CISO",        73, "Sophie Laurent","Multi-year 2yr deal closed"),
        ("D-208","Nordic Assurance",        168000,"Darktrace",  "LOST", "Darktrace AI narrative stronger",    "CEO",         90, "Sophie Laurent","Lost on AI story — needs refresh"),
        ("D-209","Quantum Payments",         48000,"CrowdStrike","WON",  "SMB pricing tier competitive",       "IT Manager",  19, "Wei Zhang",     "Quick close"),
        ("D-210","Atlas Capital Group",     360000,"SentinelOne","PENDING","POC started",                      "CISO",        41, "James Okoye",   "Strong signal — finalist"),
    ]

    for r_idx, row in enumerate(rows, 3):
        bgs = [None, None, None, None,
               GREEN if row[4] == "WON" else (RED if row[4] == "LOST" else GOLD),
               None, None, None, None, None]
        for c_idx, (val, bg) in enumerate(zip(row, bgs), 1):
            cell(ws1, c_idx, r_idx, val, bg=bg, fmt='$#,##0' if c_idx == 3 else None)

    for w, col in zip([12,20,10,14,10,28,14,10,14,28], range(1,11)):
        ws1.column_dimensions[chr(64+col)].width = w

    # ── Sheet 2: Competitor Intelligence ─────────────────────────────────
    ws2 = wb.create_sheet("Competitor_Intel")
    ws2.merge_cells("A1:F1")
    t = ws2["A1"]
    t.value = "Competitor Positioning Intel  —  Q4 FY2024"
    t.font = Font(bold=True, color="FFFFFF", size=13)
    t.fill = PatternFill("solid", fgColor=DARK)
    t.alignment = Alignment(horizontal="center", vertical="center")
    ws2.row_dimensions[1].height = 30

    for i, h in enumerate(["Competitor","Win Rate vs Us","Avg ACV ($)","Primary Win Reason","Primary Loss Reason","Battlecard Status"], 1):
        hdr(ws2, i, 2, h)

    intel = [
        ("CrowdStrike",  "41%", 285000, "Brand recognition + price on MM",   "Enterprise feature gaps (TI Add-on)",  "Current"),
        ("SentinelOne",  "38%", 310000, "AI narrative, modern UX",           "Support quality concerns post-sale",    "Needs refresh"),
        ("Darktrace",    "52%", 190000, "AI/ML story resonates with CEOs",   "Technical complexity in POC",           "Current"),
        ("Palo Alto",    "35%", 520000, "Platform breadth in Enterprise",     "Price — we win on pure endpoint play",  "Current"),
        ("Microsoft E5", "61%", 95000,  "Bundle pricing with O365",          "Dedicated product focus",               "Outdated"),
    ]
    for r_idx, row in enumerate(intel, 3):
        for c_idx, val in enumerate(row, 1):
            bg = GREEN if "Current" == val else (GOLD if "Needs refresh" == val else (RED if "Outdated" == val else None))
            cell(ws2, c_idx, r_idx, val, bg=bg if c_idx == 6 else None)

    for w, col in zip([14,12,12,28,28,14], range(1,7)):
        ws2.column_dimensions[chr(64+col)].width = w

    out = OUT / "Competitor_WinLoss_Q4FY2024.xlsx"
    wb.save(out)
    print(f"  ✓  {out.name}")


# ════════════════════════════════════════════════════════════════
# MOCK FILE 2: Customer Success QBR Notes (PDF)
# "Our CS team uses a different system — here are the QBR notes"
# Will be cross-referenced against BQ health scores
# ════════════════════════════════════════════════════════════════
def gen_pdf():
    out = OUT / "CS_QBR_Notes_Q3FY2024.pdf"
    doc = SimpleDocTemplate(str(out), pagesize=letter,
        leftMargin=0.75*inch, rightMargin=0.75*inch,
        topMargin=0.75*inch, bottomMargin=0.75*inch)

    BLUE  = HexColor("#2E5FA3")
    DARK  = HexColor("#1F3864")
    GRAY  = HexColor("#F5F5F5")
    GREEN = HexColor("#D9EAD3")
    RED   = HexColor("#FCE5CD")

    styles = getSampleStyleSheet()
    title_s  = ParagraphStyle("title", fontSize=22, textColor=DARK, spaceAfter=4, fontName="Helvetica-Bold", alignment=1)
    sub_s    = ParagraphStyle("sub",   fontSize=11, textColor=BLUE, spaceAfter=12, fontName="Helvetica", alignment=1)
    h1_s     = ParagraphStyle("h1",    fontSize=14, textColor=DARK, spaceBefore=14, spaceAfter=6, fontName="Helvetica-Bold")
    h2_s     = ParagraphStyle("h2",    fontSize=12, textColor=BLUE, spaceBefore=10, spaceAfter=4, fontName="Helvetica-Bold")
    body_s   = ParagraphStyle("body",  fontSize=10, leading=14, spaceAfter=6, fontName="Helvetica")
    note_s   = ParagraphStyle("note",  fontSize=9,  textColor=HexColor("#555555"), leading=13, spaceAfter=4, fontName="Helvetica-Oblique")

    story = []

    story.append(Spacer(1, 0.3*inch))
    story.append(Paragraph("saasmetrics.ai", title_s))
    story.append(Paragraph("Customer Success QBR Notes — Q3 FY2024", sub_s))
    story.append(Paragraph("Prepared by CS Team  |  October 1, 2024  |  INTERNAL", sub_s))
    story.append(Spacer(1, 0.2*inch))

    # Summary table
    summary_data = [
        ["Metric", "Q3 Value", "vs Q2", "Status"],
        ["Accounts with QBR completed", "11 / 14",   "+2",    "✓ On track"],
        ["Avg CSAT (post-QBR)",         "4.6 / 5",   "+0.2",  "✓ Strong"],
        ["Renewal risk accounts",        "3",         "+1",    "⚠ Monitor"],
        ["Expansion opportunities ID'd", "4",         "+1",    "✓ Good"],
        ["Executive sponsor gaps",       "2",         "same",  "⚠ Action needed"],
    ]
    t = Table(summary_data, colWidths=[2.8*inch, 1.4*inch, 1*inch, 1.5*inch])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), BLUE),
        ("TEXTCOLOR",  (0,0), (-1,0), white),
        ("FONTNAME",   (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",   (0,0), (-1,-1), 9),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [GRAY, white]),
        ("GRID",       (0,0), (-1,-1), 0.5, HexColor("#CCCCCC")),
        ("ALIGN",      (0,0), (-1,-1), "CENTER"),
        ("VALIGN",     (0,0), (-1,-1), "MIDDLE"),
        ("PADDING",    (0,0), (-1,-1), 5),
    ]))
    story.append(t)
    story.append(Spacer(1, 0.2*inch))

    # Account notes
    accounts = [
        ("Apex Financial", "HEALTHY", "QBR held Sept 12. CISO David Cheng expressed interest in expanding to 970 seats (+14%) citing new trading desk onboarding. TI Add-on renewal confirmed — strong value realization. NPS respondent: +58. Exec sponsor: David Cheng (CISO) + Mark Osei (CTO). No risk flags. Expansion deal in Q4 pipeline (D-expansion-001, $47K ACV)."),
        ("Meridian Trading", "AT RISK", "QBR attempted Sept 18 — rescheduled twice. Finally held Sept 25 via video only; incoming CTO Li Wei attended but no decision authority on renewal. CrestPeak acquisition context: IT freeze on all vendor spend >$50K pending integration review. CSM Wei Zhang has weekly check-ins. Primary risk: acquisition clause in contract allows 60-day termination notice post-close. Board notification threshold met per policy §5.3. Recommended: CRO-level engagement + 15% save discount offer before Nov 30."),
        ("Vantage Capital", "HEALTHY", "QBR held Aug 28. Smooth renewal conversation — auto-renew confirmed. 595 active users out of 620 contracted (96%). Minor ask: enhanced reporting for compliance team. Logged as feature request. Exec sponsor stable: Sarah Mills (COO)."),
        ("Castlepoint Insurance", "HEALTHY", "QBR held Sept 5. Very positive — CSAT 5/5. Seat utilization at 90% and growing (150 contracted, 135 active). Discussion of expansion to 200 seats in H1 FY25. RIIP promotion conversation initiated — qualifies as regulated industry >$75K ACV. Exec sponsor: Tom Bradley (CISO)."),
        ("Pinnacle Wealth", "AT RISK", "No QBR completed. CSM James Okoye sent 3 meeting requests over 6 weeks — no response. Seat utilization declining: 22/45 active (49%). Feature adoption at 31% — lowest in book. Last login by any user: Sept 3. Escalation recommended: executive outreach from VP CS. Contract expires March 31. Medium-high churn probability."),
        ("GoldLeaf Advisors", "HEALTHY", "First QBR (new customer Oct 1). Onboarding complete. 155/175 seats provisioned and active. Users very engaged — feature adoption 52% in first 30 days (above benchmark of 40%). CSAT 5/5. Referral conversation: they mentioned Pacific Rim Capital as a warm intro opportunity."),
        ("Pacific Mutual Bank", "HEALTHY", "QBR held Sept 10. Solid quarter. 174/190 active seats. Contract auto-renewing Oct 1 at $108K ARR. Expansion potential: compliance team requesting dedicated workspace feature — roadmap discussion ongoing. Exec sponsor: Angela Tan (CRO)."),
        ("BlueSky Fintech", "HEALTHY", "QBR held Sept 20. New customer (Aug 1). 100% seat utilization (30/30 active) after only 60 days. Extremely high engagement. NPS: +70. Potential case study candidate — flagged for marketing."),
        ("NordBank AG", "CHURNED", "Post-churn analysis only. Contract expired June 9. BaFin fines (€12M) in Q1 triggered budget freeze across all non-critical vendor spend. CISO Dieter Hoffmann departed April 14 — no exec sponsor. Renewal conversation stalled without sponsor. Classified: PREVENTABLE with earlier exec escalation. Re-engagement possible in FY26 when budget resets. POC of record: procurement contact Heike Brandt."),
    ]

    story.append(Paragraph("Account QBR Notes", h1_s))

    for acct, status, notes in accounts:
        color = GREEN if status == "HEALTHY" else (RED if status == "AT RISK" else HexColor("#EEEEEE"))
        acct_data = [[Paragraph(f"<b>{acct}</b>", body_s), Paragraph(f"<b>{status}</b>", body_s)]]
        at = Table(acct_data, colWidths=[5.5*inch, 1.2*inch])
        at.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), color),
            ("PADDING",    (0,0), (-1,-1), 6),
            ("GRID",       (0,0), (-1,-1), 0.5, HexColor("#CCCCCC")),
        ]))
        story.append(at)
        story.append(Paragraph(notes, note_s))
        story.append(Spacer(1, 0.1*inch))

    story.append(Spacer(1, 0.2*inch))
    story.append(Paragraph("Key Actions for Q4", h1_s))
    actions = [
        "Meridian Trading: CRO to engage incoming CTO Li Wei before Nov 15. Offer 15% save discount + 2-year deal.",
        "Pinnacle Wealth: VP CS executive outreach by Oct 15. If no response by Oct 31, classify as probable churn.",
        "Apex Financial: Expansion proposal (970 seats, $47K ACV) to be submitted by Wei Zhang by Oct 20.",
        "Castlepoint Insurance: RIIP promotion offer to be made at next touchpoint. Priya Nair to lead.",
        "NordBank AG: Mark for re-engagement outreach Q1 FY26. Keep Heike Brandt contact warm.",
        "GoldLeaf Advisors: Pacific Rim Capital intro to be requested by Sophie Laurent (referral potential $144K ACV).",
    ]
    for a in actions:
        story.append(Paragraph(f"• {a}", body_s))

    doc.build(story)
    print(f"  ✓  {out.name}")


# ════════════════════════════════════════════════════════════════
# MOCK FILE 3: FY2025 Budget & Headcount Plan (Word)
# "Finance sent this over — attach it to the assistant"
# Will be queried for budget context against deal sizes
# ════════════════════════════════════════════════════════════════
def gen_word():
    doc = Document()

    DARK  = RGBColor(0x1F, 0x38, 0x64)
    BLUE  = RGBColor(0x2E, 0x5F, 0xA3)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    # Title
    doc.add_heading("saasmetrics.ai", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("GTM Budget & Headcount Plan — FY2025")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = BLUE

    note = doc.add_paragraph("Finance & RevOps  |  Planning Cycle: Oct 2024  |  CONFIDENTIAL")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.runs[0].font.size = Pt(10)
    note.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    doc.add_paragraph()

    def h1(text):
        p = doc.add_heading(text, 1)
        p.runs[0].font.color.rgb = DARK
        return p

    def h2(text):
        p = doc.add_heading(text, 2)
        p.runs[0].font.color.rgb = BLUE
        return p

    def body(text):
        return doc.add_paragraph(text)

    def add_table(headers, rows, col_widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        hdr_cells = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        for row_data in rows:
            row = t.add_row().cells
            for i, val in enumerate(row_data):
                row[i].text = str(val)
                row[i].paragraphs[0].runs[0].font.size = Pt(9)
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in t.rows:
                    row.cells[i].width = Inches(w)
        return t

    h1("FY2025 GTM Revenue Targets")
    body("Targets are set against Q4 FY2024 exit ARR of $2.31M. FY2025 plan assumes $1.8M net new ARR to reach $4.1M total ARR by December 2025 — a 77% year-over-year growth target.")

    add_table(
        ["Segment", "FY2024 Exit ARR", "FY2025 Target ARR", "Net New ARR Target", "Key Assumption"],
        [
            ("Enterprise",  "$1.38M", "$2.60M", "$1.22M", "2 new logos + Meridian retention"),
            ("Mid-Market",  "$0.69M", "$1.10M", "$0.41M", "4 new logos, 5% churn max"),
            ("SMB",         "$0.24M", "$0.40M", "$0.16M", "Self-serve motion launch Q2"),
            ("TOTAL",       "$2.31M", "$4.10M", "$1.79M", "77% YoY growth"),
        ],
        [1.2, 1.1, 1.1, 1.1, 2.0],
    )

    doc.add_paragraph()
    h1("Sales Headcount Plan")
    body("Plan reflects 3 new AE hires in H1 FY2025 plus 2 SDR hires in Q1. CS team grows by 2 CSMs in Q2 to support customer base expansion.")

    add_table(
        ["Role", "Current HC", "H1 Hires", "H2 Hires", "EOY HC", "OTE ($)", "Total Cost ($)"],
        [
            ("Enterprise AE",   "2", "1", "1", "4",  "$280K", "$1.12M"),
            ("Mid-Market AE",   "2", "1", "0", "3",  "$200K", "$600K"),
            ("SMB AE",          "0", "1", "1", "2",  "$150K", "$300K"),
            ("SDR",             "1", "2", "1", "4",  "$80K",  "$320K"),
            ("CSM",             "2", "2", "0", "4",  "$130K", "$520K"),
            ("SE / Solutions",  "1", "0", "1", "2",  "$180K", "$360K"),
            ("RevOps",          "1", "0", "0", "1",  "$140K", "$140K"),
            ("TOTAL",           "9", "7", "4", "20", "—",     "$3.34M"),
        ],
        [1.3, 0.8, 0.7, 0.7, 0.7, 0.8, 0.9],
    )

    doc.add_paragraph()
    h1("Marketing Budget Allocation")
    add_table(
        ["Category", "FY2024 Budget", "FY2025 Budget", "Change", "Rationale"],
        [
            ("Demand Gen (Digital)",   "$180K", "$320K",  "+78%", "Scale paid + SEO for inbound"),
            ("Field Events / Sponsorships", "$95K", "$140K", "+47%", "RSA, Black Hat, FS-ISAC"),
            ("Content & Analyst Relations", "$60K", "$100K", "+67%", "Gartner MQ, Forrester Wave"),
            ("Partner / Channel",      "$40K",  "$80K",   "+100%","New MSSPs program launch"),
            ("Brand / Creative",       "$45K",  "$60K",   "+33%", "Refresh brand for Series B"),
            ("TOTAL",                  "$420K", "$700K",  "+67%", ""),
        ],
        [1.5, 1.0, 1.0, 0.7, 2.3],
    )

    doc.add_paragraph()
    h1("Key FY2025 Assumptions & Risks")
    assumptions = [
        ("ARR assumption", "Meridian Trading ($540K, 24% of ARR) renews at $459K (15% discount, 2-year). If churn, target missed by $540K — full year impact."),
        ("Headcount timing", "3 AE hires must be onboarded by Q1 end for full-year quota contribution. Ramp time assumed 90 days."),
        ("Series B close", "Assumes $12M Series B closes November 2024 (pre-money $48M). Budget is contingent on close."),
        ("SMB motion",      "Self-serve product launch assumed Q2. If delayed, SMB target is at risk ($160K net new ARR)."),
        ("Competition",     "CrowdStrike mid-market push could pressure win rates. Battlecard refresh scheduled Q1."),
        ("Churn cap",       "Total churn capped at $180K in the plan. Meridian save is the single biggest risk to plan."),
    ]
    for k, v in assumptions:
        p = doc.add_paragraph()
        run = p.add_run(f"{k}: ")
        run.bold = True
        run.font.size = Pt(10)
        p.add_run(v).font.size = Pt(10)

    out = OUT / "FY2025_GTM_Budget_Plan.docx"
    doc.save(out)
    print(f"  ✓  {out.name}")


# ── Run all ───────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("Generating mock upload files...")
    gen_excel()
    gen_pdf()
    gen_word()
    print(f"\nAll files written to: {OUT}/")
    print("\nDemo upload sequence:")
    print("  1. Competitor_WinLoss_Q4FY2024.xlsx  — then ask: 'Which competitors are we losing to most and why?'")
    print("  2. CS_QBR_Notes_Q3FY2024.pdf         — then ask: 'Cross-reference QBR notes with at-risk BQ data'")
    print("  3. FY2025_GTM_Budget_Plan.docx        — then ask: 'What is the FY2025 ARR target and what assumptions is it based on?'")
